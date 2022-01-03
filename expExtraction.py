from docx import Document
import re

def expExtract(docx):
    doc = Document(docx)

    nPara= len(doc.paragraphs)

    topics= ["experience", "skill", "contact", "achievement"]


    expStart=0
    expEnd=0
    skillStart=0
    skillEnd=0
    introStart=0
    introEnd=0
    # Seperating intro
    for i, para in enumerate(doc.paragraphs):
        paratext=doc.paragraphs[i].text.lower()
           
        for j in topics:
            if re.search(j,paratext)!=None:
                if j!="intro":
                    introEnd=i
                break
        if introEnd>0:
            break
    #Seperating experience
    for i, para in enumerate(doc.paragraphs):
        paratext=doc.paragraphs[i].text.lower()
        if re.search("experience",paratext)!=None:
           expStart=i
           
        for j in topics:
            if re.search(j,paratext)!=None and expStart!=0:
                if j!="experience":
                    expEnd=i
                break
        if expStart>0 and expEnd>0:
            break
        

    if expEnd==0:
        expEnd=len(doc.paragraphs)
        
    #Seperating Skills
    for i, para in enumerate(doc.paragraphs):
        paratext=doc.paragraphs[i].text.lower()
        if re.search("skill",paratext)!=None:
           skillStart=i
        for j in topics:
            if re.search(j,paratext)!=None and skillStart:
                if j!="skill":
                    skillEnd=i
        if skillStart>0 and skillEnd>0:
            break

    if skillEnd==0:
        skillEnd=len(doc.paragraphs)

# Seperating Languages
    languages=["urdu", "english", "sindhi","punjabi"]
    rawLang=""
    for i, para in enumerate(doc.paragraphs):
        paratext=doc.paragraphs[i].text.lower()
        for j in languages:
            if re.search(j,paratext)!=None:
                rawLang= rawLang+" " +j
## Seperating and performing calculations on Experience
    rawExp=""; 
    years=0
    
    for i in range(expStart, expEnd):
        txt=doc.paragraphs[i].text
        rawExp=rawExp+"\n"+txt
        x= re.findall("20..", txt)
        try:
            years = years+int(x[1]) - int(x[0])
        except:
            pass

##Seperating skills
    
    rawSkill=""
    for i in range(skillStart,skillEnd):
        txt=doc.paragraphs[i].text
        rawSkill=rawSkill+"\n"+txt

    intro=""
    for i in range(introStart,introEnd):
        txt=doc.paragraphs[i].text
        intro=intro+"\n"+txt

#raw CV for additional Keywords
    rawCV=""
    for i, para in enumerate(doc.paragraphs):
        paratext=doc.paragraphs[i].text.lower()
        rawCV= rawCV+"\n"+paratext



    return intro,years,rawExp, rawSkill, rawLang, rawCV

if __name__=="__main__":
    intro,yr,exp,ski,lang,raw=expExtract('docx1.docx')
    print(intro)
    print("--------------------------------")
    print(yr)
    print("-----------Experience----------")
    print(exp)
    print("-----------Skill----------")
    print(ski)
    print("---------Languages----------")
    print(lang)
    print("---------Raw CV------------")
    print(raw)
